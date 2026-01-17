"""Word document export service using python-docx."""

from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.project import Project
from ..models.five_whys import FiveWhys
from ..models.ishikawa import IshikawaDiagram
from ..models.action_plan import ActionPlan
from ..config.settings import settings


class WordService:
    """Service for exporting data to Word documents."""
    
    def export_project_to_word(self, project: Project, output_path: str) -> bool:
        """Export project data to a Word document.
        
        Args:
            project: Project instance to export.
            output_path: Path where Word document will be saved.
        
        Returns:
            True if successful, False otherwise.
        """
        try:
            doc = Document()
            
            # Set document properties
            doc.core_properties.author = settings.COMPANY_NAME
            doc.core_properties.title = f"{project.name} - Kaizen Blitz Report"
            
            # Title
            title = doc.add_heading(project.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            subtitle = doc.add_paragraph("Kaizen Blitz Project Report")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0]
            subtitle_format.font.size = Pt(14)
            subtitle_format.font.color.rgb = RGBColor(0, 120, 212)
            
            doc.add_paragraph()
            
            # Company and date
            company_para = doc.add_paragraph(settings.COMPANY_NAME)
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Project Overview
            doc.add_heading('Project Overview', 1)
            
            # Project details table
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Light Grid Accent 1'
            
            details = [
                ('Project Name:', project.name),
                ('Status:', project.status.value),
                ('Current Phase:', project.current_phase.value),
                ('Target Area:', project.target_area or 'N/A'),
                ('Start Date:', project.start_date.strftime('%Y-%m-%d') if project.start_date else 'N/A'),
                ('Expected Completion:', project.expected_completion_date.strftime('%Y-%m-%d') if project.expected_completion_date else 'N/A'),
                ('Progress:', f"{project.progress_percentage}%"),
            ]
            
            for i, (label, value) in enumerate(details):
                row_cells = table.rows[i].cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
                
                # Make labels bold
                row_cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Description
            if project.description:
                doc.add_heading('Description', 2)
                doc.add_paragraph(project.description)
            
            # Team Members
            team_members = project.get_team_members()
            if team_members:
                doc.add_heading('Team Members', 2)
                for member in team_members:
                    doc.add_paragraph(member, style='List Bullet')
            
            doc.add_page_break()
            
            # Five Whys sections
            for five_whys in project.five_whys:
                if five_whys.is_completed:
                    self._add_five_whys_section(doc, five_whys)
                    doc.add_page_break()
            
            # Ishikawa sections
            for ishikawa in project.ishikawa_diagrams:
                if ishikawa.is_completed:
                    self._add_ishikawa_section(doc, ishikawa)
                    doc.add_page_break()
            
            # Action plan sections
            for action_plan in project.action_plans:
                if action_plan.is_completed:
                    self._add_action_plan_section(doc, action_plan)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error exporting to Word: {e}")
            return False
    
    def _add_five_whys_section(self, doc: Document, five_whys: FiveWhys) -> None:
        """Add Five Whys section to document.
        
        Args:
            doc: Document instance.
            five_whys: FiveWhys instance.
        """
        doc.add_heading('5 Whys Analysis', 1)
        
        # Problem statement
        doc.add_heading('Problem Statement', 2)
        doc.add_paragraph(five_whys.problem_statement)
        
        # Analysis
        doc.add_heading('Analysis', 2)
        all_whys = five_whys.get_all_whys()
        
        for i, why in enumerate(all_whys, 1):
            why_para = doc.add_paragraph()
            why_para.add_run(f'Why {i}: ').bold = True
            why_para.add_run(why)
        
        # Root cause
        if five_whys.root_cause:
            doc.add_heading('Root Cause', 2)
            root_para = doc.add_paragraph(five_whys.root_cause)
            # Highlight root cause
            root_para.runs[0].font.highlight_color = 3  # Yellow highlight
    
    def _add_ishikawa_section(self, doc: Document, ishikawa: IshikawaDiagram) -> None:
        """Add Ishikawa diagram section to document.
        
        Args:
            doc: Document instance.
            ishikawa: IshikawaDiagram instance.
        """
        doc.add_heading('Ishikawa (Fishbone) Diagram', 1)
        
        # Problem statement
        doc.add_heading('Problem Statement', 2)
        doc.add_paragraph(ishikawa.problem_statement)
        
        # Categories
        doc.add_heading('Categories and Causes', 2)
        
        for category in sorted(ishikawa.categories, key=lambda c: c.order):
            doc.add_heading(category.name, 3)
            
            causes = category.get_causes_list()
            if causes:
                for cause in causes:
                    doc.add_paragraph(cause, style='List Bullet')
            else:
                doc.add_paragraph('No causes identified', style='List Bullet')
    
    def _add_action_plan_section(self, doc: Document, action_plan: ActionPlan) -> None:
        """Add action plan section to document.
        
        Args:
            doc: Document instance.
            action_plan: ActionPlan instance.
        """
        doc.add_heading('Action Plan', 1)
        
        # Completion
        completion = action_plan.calculate_completion()
        comp_para = doc.add_paragraph()
        comp_para.add_run('Completion: ').bold = True
        comp_para.add_run(f'{completion}%')
        
        doc.add_paragraph()
        
        # Tasks table
        if action_plan.tasks:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Task', 'Responsible', 'Deadline', 'Status', 'Priority', 'Notes']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Task rows
            for task in action_plan.tasks:
                row_cells = table.add_row().cells
                row_cells[0].text = task.task_description
                row_cells[1].text = task.responsible_person or 'N/A'
                row_cells[2].text = task.deadline.strftime('%Y-%m-%d') if task.deadline else 'N/A'
                row_cells[3].text = task.status.value
                row_cells[4].text = task.priority.value
                row_cells[5].text = task.notes or ''
